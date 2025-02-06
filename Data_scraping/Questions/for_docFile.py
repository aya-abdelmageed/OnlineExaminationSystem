import requests
from bs4 import BeautifulSoup
from docx import Document

# Function to scrape the questions and answers
def main(page):
    src = page.content
    soup = BeautifulSoup(src, 'lxml')
    question_div = soup.find("div", {'class': 'page-blog-text'})
    questions = question_div.find_all("h3")

    # Create a Word document
    doc = Document()
    doc.add_heading("Questions and Answers", level=1)

    # Loop through the questions and write to the Word file
    for val in questions:
        question_id = val.text
        section = val.find_all_next('p', limit=3)
        question_text = section[0].text.strip()
        choices = section[1].find_all('span')
        answer = section[2].text.strip()

        # Add question to the document
        doc.add_heading(question_id, level=2)
        doc.add_paragraph(question_text)

        # Add choices
        for choice in choices:
            doc.add_paragraph(f"Choice: {choice.get_text(strip=True)}", style='List Bullet')

        # Add the correct answer
        doc.add_paragraph(f"Answer: {answer}", style='Intense Quote')

    # Save the Word document
    doc.save("questions_and_answers.docx")

# Fetch the webpage and call the main function
page = requests.get("https://cbmceindia.com/blog/details/html-css-mcqs-with-answers/")
main(page)
